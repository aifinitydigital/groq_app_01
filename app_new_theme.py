import os
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'

import gradio as gr

from docx import Document
from datetime import datetime
from query_assistant import QueryAssistant

def download_chat(chat_history):
    """Create downloadable document from chat history"""
    doc = Document()
    doc.add_heading('Legal Consultation History', 0)
    
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Generated on: {timestamp}")
    
    for message in chat_history:
        role = "Client" if message["role"] == "user" else "Legal Assistant"
        doc.add_paragraph(f"{role}:", style='Heading 2')
        doc.add_paragraph(message["content"])
        doc.add_paragraph()
    
    doc.add_paragraph("DISCLAIMER:", style='Heading 2')
    doc.add_paragraph("This document contains general legal information and should not be construed as legal advice. Please consult with a practicing lawyer for specific legal actions.")
    
    filename = f"legal_consultation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

def clear_chat():
    """Clear chat history"""
    return "", [], {"session_id": datetime.now().strftime('%Y%m%d_%H%M%S'), "messages": []}

# Create Gradio interface

def create_interface():
    assistant = QueryAssistant()
    
    dark_theme_css = """
    body {
        background-color: #2F4F4F;
        color: #FFFFFF;
        font-family: 'sans-serif';
    }
    .gradio-container {
        background-color: #2F4F4F;
    }
    .gr-input, .gr-button, .gr-textbox {
        background-color: #3C3C3C;
        color: #FFFFFF;
        border: 1px solid #444444;
    }
    .gr-button:hover {
        background-color: #555555;
    }
    .gr-primary {
        color: #1E90FF;
    }
    """

    
    with gr.Blocks(css=dark_theme_css + "footer {visibility: hidden}") as demo:
        gr.Markdown("# Legal Assistant: Indian Law")
        gr.Markdown("""I am a legal assistant specialized in Indian Law and will help analyze your legal concerns. 
                      I provide responses in both English and Telugu.""")
        
        chatbot = gr.Chatbot(
            [],
            elem_id="chatbot",
            bubble_full_width=False,
            height="85vh",
            container=True,
            type="messages"
        )
        
        with gr.Row():
            txt = gr.Textbox(
                scale=4,
                show_label=False,
                placeholder="Describe your legal concern and press enter",
                container=False,
                submit_btn="➤"
            )
            clear_btn = gr.Button("Clear Chat")
            download_btn = gr.Button("Download Consultation")
        
        state = gr.State({"session_id": datetime.now().strftime('%Y%m%d_%H%M%S'), "messages": []})
        
        txt.submit(assistant.process_query, [txt, chatbot, state], [txt, chatbot, state])
        clear_btn.click(clear_chat, None, [txt, chatbot, state])
        download_btn.click(download_chat, inputs=[chatbot], outputs=[gr.File()])
    
    return demo

if __name__ == "__main__":
   demo = create_interface()
   demo.launch(
        server_name="0.0.0.0",
        server_port=8080,
        share=False
   )
   print("Access the interface at: http://localhost:8080")
